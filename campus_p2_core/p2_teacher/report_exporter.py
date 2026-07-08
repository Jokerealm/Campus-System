from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from campus_p2_core.contracts.p2 import P2ExamAnalysis, QuestionAnalysis


SEVERITY_LABELS = {
    "critical": "重点讲评",
    "weak": "需要巩固",
    "watch": "持续观察",
    "stable": "掌握稳定",
}


def analysis_to_markdown(analysis: P2ExamAnalysis) -> str:
    lines = [analysis.teaching_report.markdown.strip(), ""]

    lines.extend(["## 逐题分析", ""])
    for item in analysis.question_analysis:
        kp_names = "、".join(kp["name"] for kp in item.confirmed_knowledge_points) or "待确认"
        lines.append(
            f"- {item.question_no}：{SEVERITY_LABELS[item.severity]}，"
            f"得分率 {item.score_rate:.1%}，均分 {item.avg_score:g}/{item.full_score:g}，知识点：{kp_names}"
        )

    lines.extend(["", "## P3 题库检索请求", ""])
    if not analysis.p3_search_requests:
        lines.append("- 暂无需要检索的薄弱知识点。")
    for index, request in enumerate(analysis.p3_search_requests, start=1):
        codes = "、".join(request.knowledge_point_codes)
        excluded = "、".join(request.exclude_question_ids) or "无"
        lines.append(
            f"- 请求 {index}：知识点 {codes}，难度 {request.difficulty_range[0]}-"
            f"{request.difficulty_range[1]}，数量 {request.limit}，排除题目 {excluded}"
        )

    if analysis.warnings:
        lines.extend(["", "## 数据提示", ""])
        for warning in analysis.warnings:
            lines.append(f"- {warning}")

    return "\n".join(lines).strip() + "\n"


def export_analysis_markdown(analysis: P2ExamAnalysis, path: str | Path) -> Path:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(analysis_to_markdown(analysis), encoding="utf-8")
    return output_path


def export_analysis_docx(analysis: P2ExamAnalysis, path: str | Path) -> Path:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _setup_document(document)
    _add_title(document, analysis)
    _add_overview(document, analysis)
    _add_priority_questions(document, analysis.question_analysis)
    _add_knowledge_diagnostics(document, analysis)
    _add_p3_requests(document, analysis)
    _add_warnings(document, analysis)

    document.save(output_path)
    return output_path


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)


def _add_title(document: Document, analysis: P2ExamAnalysis) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(analysis.teaching_report.title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(18, 24, 38)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{analysis.class_name} | campus-system-p2 教师端诊断").font.size = Pt(10)
    document.add_paragraph("")


def _add_overview(document: Document, analysis: P2ExamAnalysis) -> None:
    _add_heading(document, "一、考试概况")
    total_full = sum(item.full_score for item in analysis.question_analysis)
    total_avg = sum(item.avg_score for item in analysis.question_analysis)
    overall_rate = total_avg / total_full if total_full else 0
    review_count = len([item for item in analysis.question_analysis if item.teacher_review_status == "pending"])

    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    headers = ["题目数", "总分", "估计均分", "整体得分率", "待教师复核"]
    values = [
        str(len(analysis.question_analysis)),
        f"{total_full:g}",
        f"{total_avg:.1f}",
        f"{overall_rate:.1%}",
        str(review_count),
    ]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
        table.cell(1, idx).text = values[idx]


def _add_priority_questions(document: Document, questions: list[QuestionAnalysis]) -> None:
    _add_heading(document, "二、优先讲评题")
    for question in sorted(questions, key=lambda item: item.score_rate)[:8]:
        kp_names = "、".join(kp["name"] for kp in question.confirmed_knowledge_points) or "待确认"
        paragraph = document.add_paragraph()
        paragraph.add_run(f"第 {question.question_no} 题").bold = True
        paragraph.add_run(
            f" | {SEVERITY_LABELS[question.severity]} | 得分率 {question.score_rate:.1%} | 知识点：{kp_names}"
        )
        if question.stem_text:
            document.add_paragraph(_clip(question.stem_text, 110))


def _add_knowledge_diagnostics(document: Document, analysis: P2ExamAnalysis) -> None:
    _add_heading(document, "三、薄弱知识点与讲评策略")
    for item in analysis.knowledge_diagnostics[:8]:
        paragraph = document.add_paragraph()
        paragraph.add_run(item.name).bold = True
        paragraph.add_run(
            f" | {SEVERITY_LABELS[item.severity]} | 得分率 {item.score_rate:.1%} | 题号 {', '.join(item.related_question_nos)}"
        )
        document.add_paragraph(item.suggestion)


def _add_p3_requests(document: Document, analysis: P2ExamAnalysis) -> None:
    _add_heading(document, "四、P3 举一反三请求")
    if not analysis.p3_search_requests:
        document.add_paragraph("暂无需要请求 P3 的薄弱知识点。")
        return
    for index, request in enumerate(analysis.p3_search_requests, start=1):
        document.add_paragraph(
            f"请求 {index}：知识点 {', '.join(request.knowledge_point_codes)}；"
            f"难度 {request.difficulty_range[0]}-{request.difficulty_range[1]}；数量 {request.limit}",
            style="List Bullet",
        )


def _add_warnings(document: Document, analysis: P2ExamAnalysis) -> None:
    if not analysis.warnings:
        return
    _add_heading(document, "五、数据提示")
    for warning in analysis.warnings:
        document.add_paragraph(warning, style="List Bullet")


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(21, 116, 110)


def _clip(text: str, max_length: int) -> str:
    if len(text) <= max_length:
        return text
    return text[: max_length - 1] + "..."
