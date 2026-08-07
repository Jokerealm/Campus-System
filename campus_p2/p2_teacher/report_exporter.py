from __future__ import annotations

import html
import re
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from campus_p2.contracts.p2 import KnowledgeDiagnostic, P2ExamAnalysis, QuestionAnalysis


REPO_ROOT = Path(__file__).resolve().parents[2]

SEVERITY_LABELS = {
    "critical": "重点讲评",
    "weak": "需要巩固",
    "watch": "持续观察",
    "stable": "掌握稳定",
}


def analysis_to_markdown(analysis: P2ExamAnalysis) -> str:
    total_full, total_avg, overall_rate, review_count = _overview_metrics(analysis)
    priority_questions = _priority_questions(analysis.question_analysis)
    weak_points = _weak_diagnostics(analysis)

    lines = [
        f"# {analysis.teaching_report.title}",
        "",
        analysis.teaching_report.summary.strip(),
        "",
        "## 课堂目标",
        "",
    ]
    for objective in _lesson_objectives(analysis, priority_questions, weak_points):
        lines.append(f"- {objective}")

    lines.extend(["", "## 诊断覆盖", ""])
    lines.extend(
        [
            f"- 题目数：{len(analysis.question_analysis)}",
            f"- 总分：{total_full:g}",
            f"- 估计均分：{total_avg:.1f}",
            f"- 整体得分率：{overall_rate:.1%}",
            f"- 知识点覆盖率：{analysis.knowledge_tag_coverage:.1%}",
            f"- 待教师复核：{review_count}",
        ]
    )

    lines.extend(["", "## 讲评安排", ""])
    for step in _lesson_flow(analysis, priority_questions, weak_points):
        lines.append(f"- {step}")

    lines.extend(["", "## 逐题分析", ""])
    for item in analysis.question_analysis:
        kp_names = "、".join(kp["name"] for kp in item.confirmed_knowledge_points) or "知识点未标注"
        lines.append(
            f"- {item.question_no}：{SEVERITY_LABELS[item.severity]}，"
            f"得分率 {item.score_rate:.1%}，均分 {item.avg_score:g}/{item.full_score:g}，知识点：{kp_names}"
        )

    lines.extend(["", "## 推荐练习题", ""])
    if not analysis.practice_recommendations:
        lines.append("- 暂无推荐练习题。")
    for group in analysis.practice_recommendations:
        if not group.items:
            lines.append(f"- {group.knowledge_point_name}：题库暂缺，建议补充校本练习。")
            continue
        lines.append(f"- {group.knowledge_point_name}：推荐 {len(group.items)} 道。")
        for index, item in enumerate(group.items[:5], start=1):
            reason = _strip_html(item.recommend_reason)
            suffix = f" 推荐理由：{reason}" if reason else ""
            lines.append(f"  - 练习 {index}：{_strip_html(item.content_html)}{suffix}")

    if analysis.warnings:
        lines.extend(["", "## 数据提示", ""])
        for warning in analysis.warnings:
            lines.append(f"- {warning}")

    return "\n".join(lines).strip() + "\n"


def export_analysis_markdown(analysis: P2ExamAnalysis, path: str | Path) -> Path:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(analysis_to_markdown(analysis), encoding="utf-8")
    return output_path


def export_analysis_docx(analysis: P2ExamAnalysis, path: str | Path) -> Path:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _setup_document(document)
    _add_title(document, analysis)
    _add_overview(document, analysis)
    _add_lesson_plan(document, analysis)
    _add_priority_questions(document, analysis.question_analysis)
    _add_knowledge_diagnostics(document, analysis)
    _add_practice_recommendations(document, analysis)
    _add_warnings(document, analysis)

    document.save(output_path)
    return output_path


def _setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_title(document: Document, analysis: P2ExamAnalysis) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(analysis.teaching_report.title or "数学考试讲评报告")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(18, 24, 38)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{analysis.class_name} | 课堂讲评与巩固建议").font.size = Pt(10)
    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary.add_run(_clip(analysis.teaching_report.summary, 120))
    summary_run.font.size = Pt(9)
    summary_run.font.color.rgb = RGBColor(86, 96, 112)
    document.add_paragraph("")


def _add_overview(document: Document, analysis: P2ExamAnalysis) -> None:
    _add_heading(document, "一、考试概况")
    total_full, total_avg, overall_rate, review_count = _overview_metrics(analysis)

    table = document.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    headers = ["题目数", "总分", "估计均分", "整体得分率", "知识点覆盖率", "待教师复核"]
    values = [
        str(len(analysis.question_analysis)),
        f"{total_full:g}",
        f"{total_avg:.1f}",
        f"{overall_rate:.1%}",
        f"{analysis.knowledge_tag_coverage:.1%}",
        str(review_count),
    ]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
        table.cell(1, idx).text = values[idx]
    _style_table_header(table)


def _add_lesson_plan(document: Document, analysis: P2ExamAnalysis) -> None:
    priority_questions = _priority_questions(analysis.question_analysis)
    weak_points = _weak_diagnostics(analysis)

    _add_heading(document, "二、课堂目标")
    for objective in _lesson_objectives(analysis, priority_questions, weak_points):
        document.add_paragraph(objective, style="List Bullet")

    _add_heading(document, "三、讲评安排")
    for step in _lesson_flow(analysis, priority_questions, weak_points):
        document.add_paragraph(step, style="List Bullet")


def _add_priority_questions(document: Document, questions: list[QuestionAnalysis]) -> None:
    _add_heading(document, "四、优先讲评题")
    for question in _priority_questions(questions):
        kp_names = "、".join(kp["name"] for kp in question.confirmed_knowledge_points) or "知识点未标注"
        paragraph = document.add_paragraph()
        paragraph.add_run(f"第 {question.question_no} 题").bold = True
        paragraph.add_run(
            f" | {SEVERITY_LABELS[question.severity]} | 得分率 {question.score_rate:.1%} | 知识点：{kp_names}"
        )
        if _question_body_text(question):
            _add_question_body(document, question)
        document.add_paragraph(_question_teaching_hint(question), style="List Bullet")


def _add_knowledge_diagnostics(document: Document, analysis: P2ExamAnalysis) -> None:
    _add_heading(document, "五、薄弱知识点与讲评策略")
    for item in _weak_diagnostics(analysis):
        paragraph = document.add_paragraph()
        paragraph.add_run(item.name).bold = True
        paragraph.add_run(
            f" | {SEVERITY_LABELS[item.severity]} | 得分率 {item.score_rate:.1%} | 题号 {', '.join(item.related_question_nos)}"
        )
        document.add_paragraph(item.suggestion)


def _add_practice_recommendations(document: Document, analysis: P2ExamAnalysis) -> None:
    _add_heading(document, "六、推荐练习题")
    if not analysis.practice_recommendations:
        document.add_paragraph("暂无推荐练习题。")
        return
    for group in analysis.practice_recommendations:
        paragraph = document.add_paragraph()
        paragraph.add_run(group.knowledge_point_name).bold = True
        paragraph.add_run(f" | 关联题号 {', '.join(group.related_question_nos) or '暂无题号'}")
        if not group.items:
            document.add_paragraph("题库暂缺同类题，建议补充校本练习。", style="List Bullet")
            continue
        for index, item in enumerate(group.items[:5], start=1):
            document.add_paragraph(f"练习 {index}：{_strip_html(item.content_html)}", style="List Bullet")
            reason = _strip_html(item.recommend_reason)
            if reason:
                document.add_paragraph(f"推荐理由：{reason}")


def _add_warnings(document: Document, analysis: P2ExamAnalysis) -> None:
    if not analysis.warnings:
        return
    _add_heading(document, "七、数据提示")
    for warning in analysis.warnings:
        document.add_paragraph(warning, style="List Bullet")


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(21, 116, 110)


def _question_body_text(question: QuestionAnalysis) -> str:
    return (getattr(question, "stem_markdown", "") or question.stem_text or "").strip()


def _add_question_body(document: Document, question: QuestionAnalysis) -> None:
    lines = _question_body_text(question).replace("\r\n", "\n").split("\n")
    has_structured_images = bool(getattr(question, "images", []))
    index = 0
    while index < len(lines):
        line = _strip_inline_image_markers(lines[index]) if has_structured_images else lines[index]
        if not line.strip():
            index += 1
            continue
        if has_structured_images and _is_markdown_image_line(line):
            index += 1
            continue
        if "|" in line:
            block: list[str] = []
            cursor = index
            while cursor < len(lines) and lines[cursor].strip() and "|" in lines[cursor]:
                block_line = _strip_inline_image_markers(lines[cursor]) if has_structured_images else lines[cursor]
                if block_line.strip() and not (has_structured_images and _is_markdown_image_line(block_line)):
                    block.append(block_line)
                cursor += 1
            if _is_likely_markdown_table(block):
                _add_markdown_table(document, block)
                index = cursor
                continue
        paragraph = document.add_paragraph()
        _add_rich_text(paragraph, line.strip())
        index += 1
    _add_question_options(document, question)
    _add_question_images(document, question)


def _strip_inline_image_markers(value: str) -> str:
    return re.sub(r"\s*\[(?:图|图片)\s*[0-9一二三四五六七八九十]+\]\s*", " ", value).strip()


def _is_markdown_image_line(value: str) -> bool:
    return bool(re.match(r"^!\[[^\]]*\]\([^)]+\)$", value.strip()))


def _add_rich_text(paragraph, text: str) -> None:
    for kind, value in _math_tokens(text):
        run = paragraph.add_run(value)
        if kind == "sup":
            run.font.superscript = True
        elif kind == "sub":
            run.font.subscript = True


def _add_question_options(document: Document, question: QuestionAnalysis) -> None:
    options = getattr(question, "options", []) or []
    for option in options:
        if not isinstance(option, dict):
            continue
        label = str(option.get("label") or "").strip()
        text = str(option.get("text") or "").strip()
        if not label and not text:
            continue
        paragraph = document.add_paragraph()
        if label:
            paragraph.add_run(f"{label}. ").bold = True
        _add_rich_text(paragraph, text)


def _math_tokens(text: str) -> list[tuple[str, str]]:
    normalized = (
        text.replace("\\times", "×")
        .replace("\\cdot", "·")
        .replace("\\leq", "≤")
        .replace("\\geq", "≥")
        .replace("\\neq", "≠")
        .replace("\\angle", "∠")
        .replace("\\triangle", "△")
        .replace("\\pi", "π")
    )
    tokens: list[tuple[str, str]] = []
    index = 0
    pattern = re.compile(r"(\\frac\{([^{}]+)\}\{([^{}]+)\}|\\sqrt\{([^{}]+)\}|\\overline\{([^{}]+)\}|\\bar\{([^{}]+)\}|\\vec\{([^{}]+)\}|\^\{([^{}]+)\}|_\{([^{}]+)\}|\^([A-Za-z0-9+\-]+)|_([A-Za-z0-9+\-]+)|_{3,})")
    for match in pattern.finditer(normalized):
        if match.start() > index:
            tokens.append(("text", normalized[index : match.start()]))
        matched = match.group(0)
        if matched.startswith("\\frac"):
            tokens.append(("text", f"({match.group(2)})/({match.group(3)})"))
        elif matched.startswith("\\sqrt"):
            tokens.append(("text", f"√{match.group(4)}"))
        elif matched.startswith("\\overline") or matched.startswith("\\bar"):
            value = match.group(5) or match.group(6) or ""
            tokens.append(("text", _combining_overline(value)))
        elif matched.startswith("\\vec"):
            tokens.append(("text", f"{match.group(7)}⃗"))
        elif matched.startswith("^{"):
            tokens.append(("sup", match.group(8)))
        elif matched.startswith("_{"):
            tokens.append(("sub", match.group(9)))
        elif matched.startswith("^"):
            tokens.append(("sup", match.group(10)))
        elif matched.startswith("_") and set(matched) == {"_"}:
            tokens.append(("text", "______"))
        elif matched.startswith("_"):
            tokens.append(("sub", match.group(11)))
        index = match.end()
    if index < len(normalized):
        tokens.append(("text", normalized[index:]))
    return tokens or [("text", normalized)]


def _combining_overline(value: str) -> str:
    return "".join(f"{char}\u0305" for char in value)


def _is_markdown_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line.strip()))


def _split_markdown_table_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|"):
        text = text[:-1]
    return [cell.strip() for cell in text.split("|")]


def _is_likely_markdown_table(block: list[str]) -> bool:
    rows = [line.strip() for line in block if line.strip()]
    if len(rows) < 2:
        return False
    split_rows = [_split_markdown_table_row(line) for line in rows if not _is_markdown_table_separator(line)]
    if not split_rows:
        return False
    enough_columns = sum(1 for row in split_rows if len(row) >= 3) >= 2
    return any(_is_markdown_table_separator(line) for line in rows) or enough_columns


def _add_markdown_table(document: Document, block: list[str]) -> None:
    rows = [_split_markdown_table_row(line) for line in block if not _is_markdown_table_separator(line)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_rich_text(paragraph, row[column_index] if column_index < len(row) else "")
    _style_table_header(table)


def _add_question_images(document: Document, question: QuestionAnalysis) -> None:
    for image in question.images or []:
        if not isinstance(image, dict):
            continue
        path = _resolve_local_image_path(str(image.get("path") or ""))
        if not path:
            continue
        try:
            document.add_picture(str(path), width=Inches(3.2))
            last = document.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            document.add_paragraph(f"题图：{path.name}")


def _resolve_local_image_path(path: str) -> Path | None:
    if not path or re.match(r"^https?://", path):
        return None
    normalized = unquote(path)
    candidates: list[Path] = []
    if normalized.startswith("/api/assets/"):
        candidates.append(REPO_ROOT / normalized.removeprefix("/api/assets/"))
    if normalized.startswith("/"):
        candidates.append(REPO_ROOT / normalized.lstrip("/"))
    candidates.append(Path(normalized))
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return None


def _clip(text: str, max_length: int) -> str:
    if len(text) <= max_length:
        return text
    return text[: max_length - 1] + "..."


def _strip_html(value: str) -> str:
    text = re.sub(r"<br\s*/?>", " ", value or "", flags=re.IGNORECASE)
    text = re.sub(r"</p\s*>", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def _overview_metrics(analysis: P2ExamAnalysis) -> tuple[float, float, float, int]:
    total_full = sum(item.full_score for item in analysis.question_analysis)
    total_avg = sum(item.avg_score for item in analysis.question_analysis)
    overall_rate = total_avg / total_full if total_full else 0
    review_count = len([item for item in analysis.question_analysis if item.teacher_review_status == "pending"])
    return total_full, total_avg, overall_rate, review_count


def _priority_questions(questions: list[QuestionAnalysis], limit: int = 8) -> list[QuestionAnalysis]:
    return sorted(questions, key=lambda item: (item.score_rate, -item.loss_rate, item.question_no))[:limit]


def _weak_diagnostics(analysis: P2ExamAnalysis, limit: int = 8) -> list[KnowledgeDiagnostic]:
    return sorted(
        analysis.knowledge_diagnostics,
        key=lambda item: (item.score_rate, -len(item.related_question_nos), item.name),
    )[:limit]


def _lesson_objectives(
    analysis: P2ExamAnalysis,
    priority_questions: list[QuestionAnalysis],
    weak_points: list[KnowledgeDiagnostic],
) -> list[str]:
    first_weak = "、".join(item.name for item in weak_points[:3]) or "本次考试暴露的薄弱知识点"
    priority_nos = "、".join(item.question_no for item in priority_questions[:5]) or "重点题"
    practice_count = sum(len(group.items[:3]) for group in analysis.practice_recommendations)
    return [
        f"围绕 {first_weak} 进行集中订正，帮助学生把错题归因落到具体知识点。",
        f"优先讲评第 {priority_nos} 题，先解决共性失分，再处理个别疑问。",
        f"完成 {practice_count or '若干'} 道同类练习的当堂巩固或课后分层布置。",
    ]


def _lesson_flow(
    analysis: P2ExamAnalysis,
    priority_questions: list[QuestionAnalysis],
    weak_points: list[KnowledgeDiagnostic],
) -> list[str]:
    weak_label = "、".join(item.name for item in weak_points[:2]) or "薄弱知识点"
    priority_label = "、".join(f"第 {item.question_no} 题" for item in priority_questions[:4]) or "重点题"
    practice_count = sum(len(group.items[:3]) for group in analysis.practice_recommendations)
    return [
        "课前 3 分钟：展示整体得分率与知识点覆盖情况，明确本节课讲评边界。",
        f"共性讲评 18 分钟：围绕 {weak_label}，串讲相关概念、易错点和解题入口。",
        f"重点题突破 15 分钟：按 {priority_label} 的顺序讲评，保留学生二次思考时间。",
        f"巩固反馈 10 分钟：布置 {practice_count or '若干'} 道同类练习，课后根据完成情况调整下一轮练习。",
    ]


def _question_teaching_hint(question: QuestionAnalysis) -> str:
    if question.score_rate < 0.45:
        return "建议先回到概念和基本方法，再展示完整解题过程。"
    if question.score_rate < 0.7:
        return "建议突出关键步骤，让学生说明易错环节。"
    return "建议快速核对方法，把时间留给同类迁移。"


def _style_table_header(table) -> None:
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
