from __future__ import annotations

from base64 import b64decode
from io import BytesIO
import json
import os
from pathlib import Path
import sys
from tempfile import TemporaryDirectory


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "backend"))

from docx import Document
from docx.shared import Inches
from fastapi.testclient import TestClient

from app.main import app
from campus_p2.p1_input.normalized_paper import validate_normalized_paper
from campus_p2.p2_teacher.service import P2TeacherService


def main() -> None:
    paper_path = ROOT / "examples" / "normalized_paper_demo.json"
    scores_path = ROOT / "examples" / "sample_exam_scores.xlsx"

    validation = validate_normalized_paper(paper_path)
    assert validation["ok"], validation
    assert validation["question_count"] == 18, validation

    result = P2TeacherService().run_analysis(
        paper_json_path=paper_path,
        score_file_path=scores_path,
        exam_id="p2_smoke_demo",
        class_name="高二3班",
        output_dir=ROOT / "data" / "exams",
    )
    analysis = result.analysis
    assert len(analysis.question_analysis) == 18
    assert len(analysis.p3_search_requests) == 8
    assert result.json_path.exists()
    assert result.markdown_path and result.markdown_path.exists()
    assert result.docx_path and result.docx_path.exists()

    client = TestClient(app)
    _assert_optional_auth_gate(client)
    for health_path in ("/health", "/api/health"):
        response = client.get(health_path)
        assert response.status_code == 200, response.text
        health = response.json()
        assert health["app"] == "campus-system-p2"
        assert health["contract"] == "paper.v0.1 + p2.v0.1"

    response = client.get("/api/p2/demo")
    assert response.status_code == 200
    payload = response.json()
    assert len(payload["question_analysis"]) == 18
    assert payload["knowledge_tag_coverage"] > 0
    assert len(payload["p3_search_requests"]) == 8
    assert sum(len(group["items"]) for group in payload["practice_recommendations"]) > 0

    response = client.get("/api/demo/readiness")
    assert response.status_code == 200, response.text
    readiness = response.json()["data"]
    assert readiness["facts"]["paper_count"] >= 5
    assert readiness["facts"]["knowledge_point_count"] >= 0
    assert readiness["facts"]["question_count"] >= 0
    assert readiness["facts"]["exam_count"] >= 0
    assert readiness["facts"]["system_test_exam_count"] >= 0
    assert readiness["facts"]["system_test_lesson_plan_count"] >= 0
    assert readiness["facts"]["system_test_practice_pack_count"] >= 0
    assert {item["key"] for item in readiness["components"]} == {"p1", "p2", "p3", "llm", "security"}
    assert readiness["security"]["file_hashing_enabled"] is True
    assert readiness["security"]["static_assets_scoped"] is True
    assert readiness["security"]["max_paper_upload_mb"] > 0
    assert isinstance(readiness["security"]["token_configured"], bool)
    assert "auth_token" not in readiness["security"]

    response = client.get(
        "/auth/session",
        headers={"X-Teacher-Id": "teacher_smoke", "X-Tenant-Id": "demo_school"},
    )
    assert response.status_code == 200, response.text
    session = response.json()["data"]
    assert session["actor"]["actor_id"] == "teacher_smoke"
    assert session["actor"]["actor_role"] == "teacher"
    assert session["actor"]["tenant_id"] == "demo_school"
    assert session["auth"]["secret_visible"] is False
    assert any(item["key"] == "exam:create" for item in session["permissions"])
    assert "token" not in session

    response = client.get("/api/p2/examples/paper")
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/json")
    assert b"paper_demo_20260704_001" in response.content

    response = client.get("/api/p2/examples/scores")
    assert response.status_code == 200
    assert len(response.content) > 1_000

    with TemporaryDirectory(prefix="campus_p1_smoke_") as temp_dir:
        mini_docx = Path(temp_dir) / "mini_math_paper.docx"
        _write_mini_docx(mini_docx)
        response = client.post(
            "/api/ai/v1/parse/paper",
            data={"stage": "junior_high", "grade": "初三"},
            files={
                "file": (
                    mini_docx.name,
                    mini_docx.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert response.status_code == 200, response.text
        payload = response.json()["data"]
        assert payload["status"] == "succeeded"
        paper_job_id = payload["job_id"]
        assert payload["summary"]["question_count"] == 2
        assert payload["paper"]["stage"] == "junior_high"
        assert payload["paper"]["questions"][0]["question_type"] == "single_choice"
        assert "2025年北京市中考数学试卷" not in payload["paper"]["questions"][-1]["stem_text"]
        assert "2025年北京市中考数学试卷" not in payload["paper"]["questions"][-1]["stem_markdown"]
        assert len(payload["source_file"]["sha256"]) == 64
        response = client.get(f"/api/ai/v1/jobs/{paper_job_id}")
        assert response.status_code == 200, response.text
        assert response.json()["data"]["job_type"] == "paper_parse"
        assert response.json()["data"]["result_url"].endswith(f"/parse/paper/{paper_job_id}/result")
        response = client.get(f"/api/ai/v1/parse/paper/{paper_job_id}/result")
        assert response.status_code == 200, response.text
        paper_result = response.json()["data"]
        assert paper_result["job_id"] == paper_job_id
        assert paper_result["questions"][0]["question_no"] == "1"
        assert paper_result["summary"]["image_count"] >= 1
        response = client.get(payload["paper_url"])
        assert response.status_code == 200, response.text
        assert response.json()["schema_version"] == "paper.v0.1"
        response = client.get("/api/assets/.env.example")
        assert response.status_code == 404

        mini_scores = Path(temp_dir) / "mini_scores.csv"
        mini_scores.write_text("题号,满分,平均分\n1,5,4\n2,5,3\n", encoding="utf-8")
        response = client.post(
            "/api/ai/v1/parse/score-excel",
            files={"file": (mini_scores.name, mini_scores.read_bytes(), "text/csv")},
        )
        assert response.status_code == 200, response.text
        score_payload = response.json()["data"]
        assert score_payload["summary"]["record_count"] == 2
        assert len(score_payload["source_file"]["sha256"]) == 64
        score_job_id = score_payload["job_id"]
        response = client.get(f"/api/ai/v1/jobs/{score_job_id}")
        assert response.status_code == 200, response.text
        assert response.json()["data"]["job_type"] == "score_excel_parse"
        assert response.json()["data"]["result_url"].endswith(f"/parse/score-excel/{score_job_id}/result")
        response = client.get(f"/api/ai/v1/parse/score-excel/{score_job_id}/result")
        assert response.status_code == 200, response.text
        score_result = response.json()["data"]
        assert score_result["job_id"] == score_job_id
        assert len(score_result["score_stats"]) == 2
        assert score_result["records"] == score_result["score_stats"]

        response = client.post(
            "/api/ai/v1/parse/score-excel",
            files={"file": ("scores.exe", b"nope", "application/octet-stream")},
        )
        assert response.status_code == 400

        response = client.post(
            "/api/ai/v1/knowledge/tag",
            json={
                "subject": "math",
                "grade": "8",
                "knowledge_version": "2026.1",
                "candidate_limit": 3,
                "questions": [
                    {
                        "question_no": "1",
                        "stem_text": "函数 f(x)=x^2 的导数是（　　）",
                        "question_type": "single_choice",
                        "options": [
                            {"label": "A", "text": "2x"},
                            {"label": "B", "text": "x"},
                        ],
                    }
                ],
            },
        )
        assert response.status_code == 200, response.text
        tag_payload = response.json()["data"]
        assert tag_payload["items"][0]["question_no"] == "1"
        assert tag_payload["items"][0]["candidates"]
        assert "knowledge_point_id" in tag_payload["items"][0]["candidates"][0]
        assert any("导数" in item["knowledge_point_name"] for item in tag_payload["items"][0]["candidates"])

        response = client.post(
            "/api/ai/v1/wrong-question/recognize",
            json={
                "student_id": "stu_smoke",
                "file": {
                    "file_id": "file_wrong_smoke",
                    "storage_uri": "local://student/stu_smoke/wrong-question.png",
                    "file_name": "triangle_wrong_question.png",
                    "mime_type": "image/png",
                    "size_bytes": 128,
                    "sha256": "demo",
                },
                "options": {"grade": "8"},
            },
        )
        assert response.status_code == 200, response.text
        wrong_job_id = response.json()["data"]["job_id"]
        assert wrong_job_id.startswith("job_wrong_")

        response = client.get(f"/api/ai/v1/jobs/{wrong_job_id}")
        assert response.status_code == 200, response.text
        job_payload = response.json()["data"]
        assert job_payload["status"] == "succeeded"
        assert job_payload["progress"] == 100
        assert job_payload["result_url"].endswith(f"/wrong-question/recognize/{wrong_job_id}/result")

        response = client.get(f"/api/ai/v1/wrong-question/recognize/{wrong_job_id}/result")
        assert response.status_code == 200, response.text
        wrong_result = response.json()["data"]
        assert wrong_result["status"] == "succeeded"
        assert wrong_result["result"]["question"]["needs_review"] is True
        assert wrong_result["result"]["knowledge_candidates"][0]["knowledge_point_id"] == "kp_math_8_triangle_side_relation"

        response = client.post(
            "/api/ai/v1/explanations/guided/next",
            json={
                "student_id": "stu_smoke",
                "wrong_question_id": "wq_smoke",
                "question_html": wrong_result["result"]["question"]["stem_html"],
                "knowledge_point_ids": ["kp_math_8_triangle_side_relation"],
                "current_step_index": 0,
                "student_input": "",
                "mode": "hint",
            },
        )
        assert response.status_code == 200, response.text
        guided_payload = response.json()["data"]
        assert guided_payload["step_index"] == 1
        assert guided_payload["content"]
        assert guided_payload["can_show_full_answer"] is False

        response = client.post(
            "/api/ai/v1/questions/variants/generate",
            json={
                "source_question_id": "wq_smoke",
                "source_question_html": wrong_result["result"]["question"]["stem_html"],
                "knowledge_point_ids": ["kp_math_8_triangle_side_relation"],
                "difficulty_target": 0.55,
                "count": 2,
                "constraints": {"question_type": "single_choice"},
            },
        )
        assert response.status_code == 200, response.text
        variant_payload = response.json()["data"]
        assert variant_payload["source"] in {"heuristic", "llm"}
        assert len(variant_payload["items"]) == 2
        assert all(item["audit_status"] == "pending_review" for item in variant_payload["items"])

        response = client.get(
            "/student/practice/recommendations?student_id=student_smoke&limit=2",
            headers={"X-Student-Id": "student_smoke", "X-Tenant-Id": "demo_school"},
        )
        assert response.status_code == 200, response.text
        student_recommendations = response.json()["data"]
        assert student_recommendations["items"], student_recommendations
        assert student_recommendations["source"] in {"p3-http", "local-fallback"}, student_recommendations
        student_question_id = student_recommendations["items"][0]["bank_question_id"]

        response = client.post(
            "/student/practice/answers",
            headers={
                "X-Student-Id": "student_smoke",
                "X-Tenant-Id": "demo_school",
                "Idempotency-Key": "p2-student-practice-smoke",
            },
            json={
                "student_id": "student_smoke",
                "bank_question_id": student_question_id,
                "answer_text": "A",
                "is_correct": True,
                "used_seconds": 12,
            },
        )
        assert response.status_code == 200, response.text
        student_answer = response.json()["data"]
        assert student_answer["answer_record_id"], student_answer
        assert student_answer["source"] in {"p3-http", "local-fallback"}, student_answer

        response = client.get(
            "/student/practice/progress?student_id=student_smoke&recent_limit=3",
            headers={"X-Student-Id": "student_smoke", "X-Tenant-Id": "demo_school"},
        )
        assert response.status_code == 200, response.text
        student_progress = response.json()["data"]
        assert student_progress["source"] in {"p3-http", "local-fallback"}, student_progress
        assert isinstance(student_progress["answer_count"], int), student_progress
        assert isinstance(student_progress["mastery"], list), student_progress
        assert isinstance(student_progress["recent_answers"], list), student_progress

        response = client.get(
            "/student/practice/history?student_id=student_smoke&limit=3",
            headers={"X-Student-Id": "student_smoke", "X-Tenant-Id": "demo_school"},
        )
        assert response.status_code == 200, response.text
        student_history = response.json()["data"]
        assert student_history["source"] in {"p3-http", "local-fallback"}, student_history
        assert isinstance(student_history["items"], list), student_history
        assert isinstance(student_history["total_count"], int), student_history

        response = client.get(
            "/student/reports/personal?student_id=student_smoke&recent_limit=3",
            headers={"X-Student-Id": "student_smoke", "X-Tenant-Id": "demo_school"},
        )
        assert response.status_code == 200, response.text
        student_report = response.json()["data"]
        assert student_report["source"] in {"p3-http", "local-fallback"}, student_report
        assert "summary" in student_report, student_report
        assert isinstance(student_report["next_actions"], list), student_report

        response = client.post(
            "/ai-generated-questions",
            json={
                "source_question_id": "wq_smoke",
                "knowledge_point_version": "2026.1",
                "model_name": variant_payload["source"],
                "prompt_version": "p2_smoke.v1",
                "items": [
                    {
                        **item,
                        "question_type": "single_choice",
                        "knowledge_point_ids": ["kp_math_8_triangle_side_relation"],
                    }
                    for item in variant_payload["items"]
                ],
            },
        )
        assert response.status_code == 200, response.text
        generated_submit = response.json()["data"]
        assert generated_submit["saved_count"] == 2
        assert generated_submit["audit_status"] == "pending_review"

        response = client.get("/ai-generated-questions?status=pending_review&limit=5")
        assert response.status_code == 200, response.text
        generated_items = response.json()["data"]["items"]
        assert generated_items
        generated_question_id = generated_items[0]["generated_question_id"]

        response = client.put(
            f"/ai-generated-questions/{generated_question_id}/review",
            json={
                "decision": "approved",
                "reviewer_id": "teacher_smoke",
                "review_comment": "smoke approved",
                "publish_to_bank": True,
            },
        )
        assert response.status_code == 200, response.text
        review_payload = response.json()["data"]
        assert review_payload["generated_question_id"] == generated_question_id
        assert review_payload["audit_status"] == "approved"

        response = client.post(
            "/exams",
            json={
                "name": "Word paper standard flow smoke",
                "subject": "math",
                "grade": "junior_high",
                "class_ids": ["class_demo"],
                "exam_date": "2026-07-30",
                "teacher_id": "teacher_demo",
            },
        )
        assert response.status_code == 200, response.text
        word_exam_id = response.json()["data"]["exam_id"]

        response = client.post(
            f"/exams/{word_exam_id}/files",
            data={"file_type": "paper"},
            files={
                "file": (
                    mini_docx.name,
                    mini_docx.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert response.status_code == 200, response.text
        word_paper_file_id = response.json()["data"]["file"]["file_id"]
        assert len(response.json()["data"]["file"]["sha256"]) == 64

        response = client.post(
            f"/exams/{word_exam_id}/files",
            data={"file_type": "paper"},
            files={"file": ("paper.exe", b"nope", "application/octet-stream")},
        )
        assert response.status_code == 400

        response = client.post(
            f"/exams/{word_exam_id}/files",
            data={"file_type": "score_excel"},
            files={"file": (mini_scores.name, mini_scores.read_bytes(), "text/csv")},
        )
        assert response.status_code == 200, response.text
        word_score_file_id = response.json()["data"]["file"]["file_id"]
        assert len(response.json()["data"]["file"]["sha256"]) == 64

        response = client.post(
            f"/exams/{word_exam_id}/parse",
            json={
                "score_file_id": word_score_file_id,
                "paper_file_id": word_paper_file_id,
                "auto_tag_knowledge": True,
            },
        )
        assert response.status_code == 200, response.text
        assert response.json()["data"]["status"] == "teacher_review"
        structure = client.get(f"/exams/{word_exam_id}/structure").json()["data"]
        assert len(structure["questions"]) == 2
        image_questions = [item for item in structure["questions"] if item["images"]]
        assert image_questions, structure
        image_path = image_questions[0]["images"][0]["path"]
        assert image_path.startswith("/api/assets/data/p1_parse_outputs/"), image_path
        response = client.get(image_path)
        assert response.status_code == 200, response.text
        assert response.headers["content-type"].startswith("image/"), response.headers

        response = client.put(
            f"/exams/{word_exam_id}/questions/{image_questions[0]['exam_question_id']}",
            json={
                "question_no": image_questions[0]["question_no"],
                "stem_html": image_questions[0]["stem_html"],
                "question_type": image_questions[0]["question_type"],
                "full_score": image_questions[0]["full_score"],
                "images": [{"image_id": "teacher_img_01", "path": image_path, "role": "stem"}],
            },
        )
        assert response.status_code == 200, response.text
        refreshed = client.get(f"/exams/{word_exam_id}/analysis").json()["data"]
        assert any(
            question["images"] and question["images"][0]["image_id"] == "teacher_img_01"
            for question in refreshed["question_analysis"]
        ), refreshed["question_analysis"]

    with paper_path.open("rb") as paper_file, scores_path.open("rb") as score_file:
        response = client.post(
            "/api/p2/analyze",
            data={"exam_id": "api_smoke_demo", "class_name": "高二3班"},
            files={
                "paper_file": ("normalized_paper_demo.json", paper_file, "application/json"),
                "score_file": (
                    "sample_exam_scores.xlsx",
                    score_file,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                ),
            },
        )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert len(payload["question_analysis"]) == 18
    assert len(payload["p3_search_requests"]) == 8
    assert sum(len(group["items"]) for group in payload["practice_recommendations"]) > 0

    response = client.post("/api/p2/reports/docx", json=payload)
    assert response.status_code == 200
    assert len(response.content) > 10_000
    docx_text = _docx_text(response.content)
    assert "课堂目标" in docx_text
    assert "讲评安排" in docx_text
    assert "练习 1" in docx_text
    assert "paper_demo_" not in docx_text

    response = client.post("/api/p2/reports/markdown", json=payload)
    assert response.status_code == 200
    assert "课堂目标" in response.text
    assert "讲评安排" in response.text
    assert "逐题分析" in response.text
    assert "知识点覆盖率" in response.text
    assert "P3 题库检索请求" not in response.text
    assert "paper_demo_" not in response.text

    response = client.post(
        "/exams",
        json={
            "name": "Demo math exam smoke",
            "subject": "math",
            "grade": "senior_high",
            "class_ids": ["class_demo"],
            "exam_date": "2026-07-08",
            "teacher_id": "teacher_demo",
            "is_system_test": True,
        },
    )
    assert response.status_code == 200, response.text
    exam_id = response.json()["data"]["exam_id"]

    with scores_path.open("rb") as score_file:
        response = client.post(
            f"/exams/{exam_id}/files",
            data={"file_type": "score_excel"},
            files={
                "file": (
                    "sample_exam_scores.xlsx",
                    score_file,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            },
        )
    assert response.status_code == 200, response.text
    score_file_id = response.json()["data"]["file"]["file_id"]

    with paper_path.open("rb") as paper_file:
        response = client.post(
            f"/exams/{exam_id}/files",
            data={"file_type": "paper"},
            files={"file": ("normalized_paper_demo.json", paper_file, "application/json")},
        )
    assert response.status_code == 200, response.text
    paper_file_id = response.json()["data"]["file"]["file_id"]

    response = client.post(
        f"/exams/{exam_id}/parse",
        json={
            "score_file_id": score_file_id,
            "paper_file_id": paper_file_id,
            "auto_tag_knowledge": True,
        },
    )
    assert response.status_code == 200, response.text
    assert response.json()["data"]["status"] == "teacher_review"

    response = client.get(f"/exams/{exam_id}/structure")
    assert response.status_code == 200, response.text
    structure = response.json()["data"]
    assert len(structure["questions"]) == 18
    first_question = structure["questions"][0]

    response = client.get(f"/exams/{exam_id}/analysis")
    assert response.status_code == 200, response.text
    standard_analysis = response.json()["data"]
    assert len(standard_analysis["question_analysis"]) == 18
    assert standard_analysis["knowledge_tag_coverage"] > 0
    assert standard_analysis["class_name"] == "class_demo"

    response = client.put(
        f"/exams/{exam_id}/questions/{first_question['exam_question_id']}",
        json={
            "question_no": first_question["question_no"],
            "stem_html": "<p>teacher revised stem for persistence smoke</p>",
            "question_type": first_question["question_type"] or "single_choice",
            "full_score": first_question["full_score"],
        },
    )
    assert response.status_code == 200, response.text
    assert response.json()["data"]["updated"] is True

    knowledge_ids = first_question["knowledge_point_ids"] or ["kp_teacher_smoke"]
    response = client.put(
        f"/exams/{exam_id}/questions/{first_question['exam_question_id']}/knowledge-tags",
        json={"knowledge_point_ids": knowledge_ids, "comment": "smoke confirmation"},
    )
    assert response.status_code == 200, response.text
    assert response.json()["data"]["knowledge_point_ids"] == knowledge_ids

    response = client.get(f"/exams/{exam_id}/analysis")
    assert response.status_code == 200, response.text
    standard_analysis = response.json()["data"]
    assert any(
        item["stem_text"] == "teacher revised stem for persistence smoke"
        and item["teacher_review_status"] == "confirmed"
        for item in standard_analysis["question_analysis"]
    )

    response = client.get("/exams?include_system=true")
    assert response.status_code == 200, response.text
    assert any(item["exam_id"] == exam_id for item in response.json()["data"]["items"])
    response = client.get("/exams")
    assert response.status_code == 200, response.text
    assert all(item["exam_id"] != exam_id for item in response.json()["data"]["items"])

    response = client.post(
        f"/exams/{exam_id}/diagnostics/run",
        json={
            "analysis_scope": "class",
            "class_id": "class_demo",
            "include_teaching_suggestions": True,
            "include_question_recommendations": True,
        },
    )
    assert response.status_code == 200, response.text
    diagnostic_id = response.json()["data"]["diagnostic_id"]

    response = client.get(f"/exams/{exam_id}/diagnostics/{diagnostic_id}")
    assert response.status_code == 200, response.text
    assert response.json()["data"]["summary"]["question_count"] == 18
    assert response.json()["data"]["summary"]["knowledge_tag_coverage"] > 0

    response = client.post(
        f"/exams/{exam_id}/practice-packs",
        json={
            "diagnostic_id": diagnostic_id,
            "title": "Smoke 薄弱知识点训练包",
            "target": "class",
            "target_ref_id": "class_demo",
            "created_by": "teacher_demo",
        },
    )
    assert response.status_code == 200, response.text
    practice_pack = response.json()["data"]["practice_pack"]
    assert practice_pack["practice_pack_id"].startswith("pack_")
    assert practice_pack["question_ids"]

    response = client.get(f"/exams/{exam_id}/practice-packs")
    assert response.status_code == 200, response.text
    practice_pack_items = response.json()["data"]["items"]
    assert len(practice_pack_items) >= 1
    assert practice_pack_items[0]["practice_pack_id"].startswith("pack_")

    response = client.post(
        f"/exams/{exam_id}/lesson-plans",
        json={
            "diagnostic_id": diagnostic_id,
            "template_id": "tpl_school_math_review_v1",
            "sections": ["exam_summary", "high_loss_questions", "weakness_summary"],
        },
    )
    assert response.status_code == 200, response.text
    lesson_payload = response.json()["data"]
    assert lesson_payload["file"]["file_name"].endswith(".docx")
    assert lesson_payload["file"]["download_url"].endswith("/download")

    response = client.get(f"/exams/{exam_id}/lesson-plans")
    assert response.status_code == 200, response.text
    lesson_items = response.json()["data"]["items"]
    assert len(lesson_items) >= 1
    assert lesson_items[0]["download_url"].endswith("/download")

    response = client.get(lesson_payload["file"]["download_url"])
    assert response.status_code == 200, response.text
    assert len(response.content) > 10_000
    lesson_text = _docx_text(response.content)
    assert "课堂目标" in lesson_text
    assert "讲评安排" in lesson_text
    assert "课堂讲评与巩固建议" in lesson_text
    assert "paper_demo_" not in lesson_text

    response = client.get(f"/exams/{exam_id}")
    assert response.status_code == 200, response.text
    assert response.json()["data"]["lesson_plan_count"] >= 1
    assert response.json()["data"]["latest_lesson_plan"]["download_url"].endswith("/download")
    assert response.json()["data"]["practice_pack_count"] >= 1
    assert response.json()["data"]["latest_practice_pack"]["practice_pack_id"].startswith("pack_")

    response = client.get("/audit-logs?limit=100")
    assert response.status_code == 200, response.text
    audit_items = response.json()["data"]["items"]
    audit_events = {item["event"] for item in audit_items}
    assert "exam_parse_succeeded" in audit_events
    assert "knowledge_tags_confirmed" in audit_events
    assert "diagnostic_run" in audit_events
    assert "ai_generated_question_reviewed" in audit_events
    assert "lesson_plan_downloaded" in audit_events

    print("campus-system-p2 smoke test passed")
    print(f"questions={len(analysis.question_analysis)}")
    print(f"knowledge_tag_coverage={analysis.knowledge_tag_coverage}")
    print(f"p3_requests={len(analysis.p3_search_requests)}")
    print(f"json={result.json_path}")
    print(f"markdown={result.markdown_path}")
    print(f"docx={result.docx_path}")


def _write_mini_docx(path: Path) -> None:
    image_path = path.with_suffix(".png")
    image_path.write_bytes(
        b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    document = Document()
    for index, text in enumerate(
        [
        "一、选择题",
        "1．（5分）函数 f(x)=x^2 的导数是（　　）",
        "A．2x",
        "B．x",
        "C．1",
        "D．0",
        "二、填空题",
        "2．（5分）若 x+1=3，则 x=____．",
        "2025年北京市中考数学试卷",
        "参考答案",
        "1．A",
        "2．2",
        ],
    ):
        document.add_paragraph(text)
        if index == 1:
            document.add_picture(str(image_path), width=Inches(0.5))
    document.save(path)


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(texts)


def _assert_optional_auth_gate(client: TestClient) -> None:
    previous_auth_required = os.environ.get("CAMPUS_AUTH_REQUIRED")
    previous_auth_token = os.environ.get("CAMPUS_P2_AUTH_TOKEN")
    previous_identity_required = os.environ.get("CAMPUS_IDENTITY_DIRECTORY_REQUIRED")
    previous_identity_path = os.environ.get("CAMPUS_IDENTITY_DIRECTORY_PATH")
    try:
        os.environ["CAMPUS_AUTH_REQUIRED"] = "true"
        os.environ["CAMPUS_P2_AUTH_TOKEN"] = "p2-smoke-token"
        os.environ.pop("CAMPUS_IDENTITY_DIRECTORY_REQUIRED", None)
        os.environ.pop("CAMPUS_IDENTITY_DIRECTORY_PATH", None)
        response = client.get("/exams")
        assert response.status_code == 401, response.text
        response = client.get("/auth/session")
        assert response.status_code == 401, response.text
        response = client.get(
            "/exams",
            headers={
                "Authorization": "Bearer p2-smoke-token",
                "X-Teacher-Id": "teacher_smoke",
                "X-Tenant-Id": "tenant_smoke",
            },
        )
        assert response.status_code == 200, response.text
        response = client.get(
            "/exams",
            headers={
                "Authorization": "Bearer p2-smoke-token",
                "X-Student-Id": "student_smoke",
                "X-Tenant-Id": "tenant_smoke",
                "X-Client-Role": "teacher",
            },
        )
        assert response.status_code == 403, response.text
        response = client.get(
            "/auth/session",
            headers={
                "Authorization": "Bearer p2-smoke-token",
                "X-Student-Id": "student_smoke",
                "X-Tenant-Id": "tenant_smoke",
                "X-Client-Role": "teacher",
            },
        )
        assert response.status_code == 200, response.text
        assert response.json()["data"]["actor"]["actor_role"] == "student"
        response = client.post(
            "/exams",
            json={"name": "越权考试创建", "subject": "math", "grade": "高一"},
            headers={
                "Authorization": "Bearer p2-smoke-token",
                "X-Student-Id": "student_smoke",
                "X-Tenant-Id": "tenant_smoke",
            },
        )
        assert response.status_code == 403, response.text
        response = client.get(
            "/auth/session",
            headers={
                "Authorization": "Bearer p2-smoke-token",
                "X-Teacher-Id": "teacher_smoke",
                "X-Tenant-Id": "tenant_smoke",
            },
        )
        assert response.status_code == 200, response.text
        session = response.json()["data"]
        assert session["auth"]["mode"] == "bearer_token"
        assert session["auth"]["token_configured"] is True
        assert session["actor"]["actor_role"] == "teacher"
        assert session["actor"]["tenant_id"] == "tenant_smoke"

        with TemporaryDirectory(prefix="campus_identity_smoke_") as temp_dir:
            directory_path = Path(temp_dir) / "identity_directory.json"
            directory_path.write_text(
                json.dumps(
                    {
                        "tenants": [{"tenant_id": "tenant_smoke", "name": "Smoke School"}],
                        "users": [
                            {
                                "actor_id": "teacher_smoke",
                                "role": "teacher",
                                "tenant_id": "tenant_smoke",
                                "display_name": "Smoke Teacher",
                                "active": True,
                            },
                            {
                                "actor_id": "student_smoke",
                                "role": "student",
                                "tenant_id": "tenant_smoke",
                                "display_name": "Smoke Student",
                                "active": True,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            os.environ["CAMPUS_IDENTITY_DIRECTORY_REQUIRED"] = "true"
            os.environ["CAMPUS_IDENTITY_DIRECTORY_PATH"] = str(directory_path)

            response = client.get(
                "/auth/session",
                headers={
                    "Authorization": "Bearer p2-smoke-token",
                    "X-Teacher-Id": "teacher_smoke",
                    "X-Tenant-Id": "tenant_smoke",
                },
            )
            assert response.status_code == 200, response.text
            directory_session = response.json()["data"]
            assert directory_session["actor"]["display_name"] == "Smoke Teacher"
            assert directory_session["actor"]["identity_source"] == "configured_directory"
            assert directory_session["auth"]["identity_directory_required"] is True
            assert directory_session["identity_directory"]["configured"] is True
            assert directory_session["identity_directory"]["user_count"] == 2

            response = client.get(
                "/auth/session",
                headers={
                    "Authorization": "Bearer p2-smoke-token",
                    "X-Teacher-Id": "unknown_teacher",
                    "X-Tenant-Id": "tenant_smoke",
                },
            )
            assert response.status_code == 403, response.text

            response = client.get(
                "/exams",
                headers={
                    "Authorization": "Bearer p2-smoke-token",
                    "X-Student-Id": "student_smoke",
                    "X-Tenant-Id": "tenant_smoke",
                },
            )
            assert response.status_code == 403, response.text
    finally:
        if previous_auth_required is None:
            os.environ.pop("CAMPUS_AUTH_REQUIRED", None)
        else:
            os.environ["CAMPUS_AUTH_REQUIRED"] = previous_auth_required
        if previous_auth_token is None:
            os.environ.pop("CAMPUS_P2_AUTH_TOKEN", None)
        else:
            os.environ["CAMPUS_P2_AUTH_TOKEN"] = previous_auth_token
        if previous_identity_required is None:
            os.environ.pop("CAMPUS_IDENTITY_DIRECTORY_REQUIRED", None)
        else:
            os.environ["CAMPUS_IDENTITY_DIRECTORY_REQUIRED"] = previous_identity_required
        if previous_identity_path is None:
            os.environ.pop("CAMPUS_IDENTITY_DIRECTORY_PATH", None)
        else:
            os.environ["CAMPUS_IDENTITY_DIRECTORY_PATH"] = previous_identity_path


if __name__ == "__main__":
    main()
