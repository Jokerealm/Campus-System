from __future__ import annotations

import argparse
import json
from pathlib import Path
from tempfile import TemporaryDirectory
from urllib import error, parse, request
from uuid import uuid4

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PUBLISH_GENERATED_SMOKE_ID = "genq_http_smoke_publish_once"


def main() -> None:
    parser = argparse.ArgumentParser(description="Smoke test running Campus-System HTTP services.")
    parser.add_argument("--p2-url", default="http://127.0.0.1:8000")
    parser.add_argument("--p3-url", default="http://127.0.0.1:8103")
    parser.add_argument("--frontend-url", default="http://127.0.0.1:5176")
    parser.add_argument("--min-papers", type=int, default=50)
    parser.add_argument("--min-questions", type=int, default=1200)
    parser.add_argument("--min-knowledge-points", type=int, default=300)
    parser.add_argument("--timeout", type=float, default=12)
    parser.add_argument("--paper-example", default=str(ROOT / "examples" / "normalized_paper_demo.json"))
    parser.add_argument("--scores-example", default=str(ROOT / "examples" / "sample_exam_scores.xlsx"))
    args = parser.parse_args()

    p2_base = args.p2_url.rstrip("/")
    p3_base = args.p3_url.rstrip("/")
    frontend_base = args.frontend_url.rstrip("/")

    for health_path in ("/health", "/api/health"):
        p2_health = get_json(f"{p2_base}{health_path}", timeout=args.timeout)
        assert p2_health.get("app") == "campus-system-p2", p2_health
        assert p2_health.get("contract") == "paper.v0.1 + p2.v0.1", p2_health
    expect_http_status(f"{p2_base}/api/assets/.env.example", 404, timeout=args.timeout)

    p3_health = get_json(f"{p3_base}/api/health/", timeout=args.timeout)
    assert envelope_data(p3_health).get("status") == "ok", p3_health

    stats = envelope_data(
        get_json(
            f"{p3_base}/api/resource/v1/stats?version=2026.1",
            headers={"X-Service-Id": "p2-service"},
            timeout=args.timeout,
        )
    )
    assert stats["knowledge_point_count"] >= args.min_knowledge_points, stats
    assert stats["approved_question_count"] >= args.min_questions, stats

    p3_search = envelope_data(
        post_json(
            f"{p3_base}/api/resource/v1/questions/search",
            {
                "knowledge_point_ids": ["kp_math_junior_statistics"],
                "knowledge_point_version": "2026.1",
                "difficulty_range": [0.35, 0.75],
                "source_priority": ["middle_exam_real", "school_bank", "exam_history"],
                "limit": 5,
            },
            headers={"X-Service-Id": "p2-service"},
            timeout=args.timeout,
        )
    )
    assert len(p3_search["items"]) >= 3, p3_search

    readiness = envelope_data(get_json(f"{p2_base}/api/demo/readiness", timeout=args.timeout))
    assert readiness["facts"]["paper_count"] >= args.min_papers, readiness
    assert readiness["facts"]["question_count"] >= args.min_questions, readiness
    assert readiness["facts"]["knowledge_point_count"] >= args.min_knowledge_points, readiness
    assert readiness["p3"]["connected"] is True, readiness
    assert {item["key"] for item in readiness["components"]} == {"p1", "p2", "p3", "llm", "security"}
    assert readiness["security"]["file_hashing_enabled"] is True, readiness
    assert readiness["security"]["static_assets_scoped"] is True, readiness
    assert readiness["security"]["max_paper_upload_mb"] > 0, readiness
    assert "auth_token" not in readiness["security"], readiness
    initial_customer_exam_count = readiness["facts"].get("exam_count", 0)
    initial_customer_lesson_plan_count = readiness["facts"].get("lesson_plan_count", 0)
    initial_customer_practice_pack_count = readiness["facts"].get("practice_pack_count", 0)
    initial_system_test_exam_count = readiness["facts"].get("system_test_exam_count", 0)

    teacher_session = envelope_data(
        get_json(
            f"{p2_base}/auth/session",
            headers={"X-Teacher-Id": "teacher_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert teacher_session["actor"]["actor_role"] == "teacher", teacher_session
    assert teacher_session["actor"]["actor_id"] == "teacher_http_smoke", teacher_session
    assert teacher_session["auth"]["secret_visible"] is False, teacher_session
    assert teacher_session["identity_directory"]["source"] in {"builtin_demo", "configured_directory"}, teacher_session
    assert isinstance(teacher_session["identity_directory"]["user_count"], int), teacher_session
    assert any(item["key"] == "exam:create" for item in teacher_session["permissions"]), teacher_session

    student_session = envelope_data(
        get_json(
            f"{p2_base}/auth/session",
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_session["actor"]["actor_role"] == "student", student_session
    assert student_session["actor"]["actor_id"] == "student_http_smoke", student_session
    assert student_session["auth"]["secret_visible"] is False, student_session
    assert any(item["key"] == "practice-history:read" for item in student_session["permissions"]), student_session

    wrong_job = envelope_data(
        post_json(
            f"{p2_base}/api/ai/v1/wrong-question/recognize",
            {
                "student_id": "student_http_smoke",
                "file": {
                    "file_id": "file_http_smoke",
                    "storage_uri": "local://student_http_smoke/wrong-question.png",
                    "file_name": "triangle_wrong_question.png",
                    "mime_type": "image/png",
                    "size_bytes": 128,
                    "sha256": "demo",
                },
                "options": {"subject": "math", "grade": "8"},
            },
            timeout=args.timeout,
        )
    )
    wrong_job_id = wrong_job["job_id"]
    job_status = envelope_data(get_json(f"{p2_base}/api/ai/v1/jobs/{wrong_job_id}", timeout=args.timeout))
    assert job_status["status"] == "succeeded", job_status

    wrong_result = envelope_data(
        get_json(f"{p2_base}/api/ai/v1/wrong-question/recognize/{wrong_job_id}/result", timeout=args.timeout)
    )
    assert wrong_result["status"] == "succeeded", wrong_result
    question = wrong_result["result"]["question"]
    candidates = wrong_result["result"]["knowledge_candidates"]
    assert candidates[0]["knowledge_point_id"] == "kp_math_8_triangle_side_relation", candidates

    guided = envelope_data(
        post_json(
            f"{p2_base}/api/ai/v1/explanations/guided/next",
            {
                "student_id": "student_http_smoke",
                "wrong_question_id": wrong_job_id,
                "question_html": question["stem_html"],
                "knowledge_point_ids": ["kp_math_8_triangle_side_relation"],
                "current_step_index": 0,
                "student_input": "",
                "mode": "hint",
            },
            timeout=args.timeout,
        )
    )
    assert guided["step_index"] == 1 and guided["content"], guided

    variants = envelope_data(
        post_json(
            f"{p2_base}/api/ai/v1/questions/variants/generate",
            {
                "source_question_id": wrong_job_id,
                "source_question_html": question["stem_html"],
                "knowledge_point_ids": ["kp_math_8_triangle_side_relation"],
                "difficulty_target": 0.55,
                "count": 2,
                "constraints": {"question_type": "single_choice"},
            },
            timeout=args.timeout,
        )
    )
    assert len(variants["items"]) == 2, variants

    student_recommendations = envelope_data(
        get_json(
            f"{p2_base}/student/practice/recommendations?student_id=student_http_smoke&limit=3",
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_recommendations["source"] == "p3-http", student_recommendations
    assert len(student_recommendations["items"]) >= 1, student_recommendations
    student_question_id = student_recommendations["items"][0]["bank_question_id"]
    student_answer = envelope_data(
        post_json(
            f"{p2_base}/student/practice/answers",
            {
                "student_id": "student_http_smoke",
                "bank_question_id": student_question_id,
                "answer_text": "A",
                "is_correct": True,
                "used_seconds": 18,
            },
            headers={
                "X-Student-Id": "student_http_smoke",
                "X-Tenant-Id": "demo_school",
                "Idempotency-Key": f"student-http-smoke-{uuid4().hex}",
            },
            timeout=args.timeout,
        )
    )
    assert student_answer["source"] == "p3-http", student_answer
    assert student_answer["answer_record_id"], student_answer
    assert isinstance(student_answer["updated_mastery"], list), student_answer

    student_progress = envelope_data(
        get_json(
            f"{p2_base}/student/practice/progress?student_id=student_http_smoke&recent_limit=3",
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_progress["source"] == "p3-http", student_progress
    assert student_progress["answer_count"] >= 1, student_progress
    assert student_progress["correct_count"] >= 1, student_progress
    assert isinstance(student_progress["mastery"], list), student_progress
    assert isinstance(student_progress["recent_answers"], list), student_progress

    student_history = envelope_data(
        get_json(
            f"{p2_base}/student/practice/history?student_id=student_http_smoke&limit=3",
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_history["source"] == "p3-http", student_history
    assert student_history["total_count"] >= 1, student_history
    assert len(student_history["items"]) >= 1, student_history
    assert "answer_html" not in student_history["items"][0], student_history
    assert "analysis_html" not in student_history["items"][0], student_history

    student_report = envelope_data(
        get_json(
            f"{p2_base}/student/reports/personal?student_id=student_http_smoke&recent_limit=3",
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_report["source"] == "p3-http", student_report
    assert student_report["summary"]["answer_count"] >= 1, student_report
    assert isinstance(student_report["mastery"]["weak"], list), student_report
    assert isinstance(student_report["next_actions"], list), student_report

    student_wrong_upload = envelope_data(
        post_multipart(
            f"{p2_base}/student/wrong-questions",
            fields={"student_id": "student_http_smoke", "subject": "math", "grade": "8"},
            files=[("image", "wrong-question.png", b"\x89PNG\r\n\x1a\nmock-image", "image/png")],
            headers={
                "X-Student-Id": "student_http_smoke",
                "X-Tenant-Id": "demo_school",
                "Idempotency-Key": f"student-wrong-http-smoke-{uuid4().hex}",
            },
            timeout=args.timeout,
        )
    )
    assert student_wrong_upload["source"] == "p3-http", student_wrong_upload
    assert student_wrong_upload["wrong_question_id"], student_wrong_upload
    student_wrong_detail = envelope_data(
        get_json(
            f"{p2_base}/student/wrong-questions/{student_wrong_upload['wrong_question_id']}?student_id=student_http_smoke",
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_wrong_detail["source"] == "p3-http", student_wrong_detail
    assert student_wrong_detail["status"] == "recognized", student_wrong_detail
    candidate_ids = [
        item["knowledge_point_id"]
        for item in student_wrong_detail["knowledge_candidates"]
        if item.get("knowledge_point_id")
    ]
    assert candidate_ids, student_wrong_detail
    student_wrong_confirm = envelope_data(
        put_json(
            f"{p2_base}/student/wrong-questions/{student_wrong_upload['wrong_question_id']}/confirm?student_id=student_http_smoke",
            {
                "stem_html": student_wrong_detail["question"]["stem_html"],
                "question_type": student_wrong_detail["question"].get("question_type", "选择题"),
                "knowledge_point_ids": candidate_ids,
            },
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_wrong_confirm["status"] == "confirmed", student_wrong_confirm
    student_wrong_explanation = envelope_data(
        post_json(
            f"{p2_base}/student/wrong-questions/{student_wrong_upload['wrong_question_id']}/explanation/next?student_id=student_http_smoke",
            {"current_step_index": 0, "student_input": "", "mode": "hint"},
            headers={"X-Student-Id": "student_http_smoke", "X-Tenant-Id": "demo_school"},
            timeout=args.timeout,
        )
    )
    assert student_wrong_explanation["source"] == "p3-http", student_wrong_explanation
    assert student_wrong_explanation["content"], student_wrong_explanation

    publish_review = ensure_publish_smoke_generated_question(
        p2_base=p2_base,
        p3_base=p3_base,
        source_question_id=wrong_job_id,
        variant=variants["items"][0],
        model_name=variants["source"],
        timeout=args.timeout,
    )
    assert publish_review["audit_status"] == "approved", publish_review
    assert publish_review.get("bank_question_id"), publish_review

    review_smoke_id = f"genq_http_smoke_review_{uuid4().hex[:10]}"
    generated_submit = envelope_data(
        post_json(
            f"{p2_base}/ai-generated-questions",
            generated_question_import_payload(
                source_question_id=wrong_job_id,
                model_name=variants["source"],
                items=[generated_question_item(variants["items"][-1], review_smoke_id)],
            ),
            timeout=args.timeout,
        )
    )
    assert generated_submit["saved_count"] == 1, generated_submit
    assert generated_submit["audit_status"] == "pending_review", generated_submit
    assert generated_submit["items"][0]["generated_question_id"] == review_smoke_id, generated_submit

    pending_generated = envelope_data(
        get_json(f"{p2_base}/ai-generated-questions?status=pending_review&limit=100", timeout=args.timeout)
    )
    assert any(item["generated_question_id"] == review_smoke_id for item in pending_generated["items"]), (
        pending_generated
    )

    generated_review = envelope_data(
        put_json(
            f"{p2_base}/ai-generated-questions/{review_smoke_id}/review",
            {
                "decision": "approved",
                "reviewer_id": "teacher_http_smoke",
                "review_comment": "HTTP smoke approved without publishing to the customer question bank.",
                "publish_to_bank": False,
            },
            timeout=args.timeout,
        )
    )
    assert generated_review["generated_question_id"] == review_smoke_id, generated_review
    assert generated_review["audit_status"] == "approved", generated_review
    assert generated_review.get("bank_question_id") is None, generated_review

    paper_path = Path(args.paper_example)
    scores_path = Path(args.scores_example)
    assert paper_path.exists(), paper_path
    assert scores_path.exists(), scores_path

    with TemporaryDirectory(prefix="campus_p1_http_smoke_") as temp_dir:
        p1_docx = Path(temp_dir) / "http_smoke_word_paper.docx"
        write_smoke_docx(p1_docx)
        p1_paper = envelope_data(
            post_multipart(
                f"{p2_base}/api/ai/v1/parse/paper",
                fields={"stage": "junior_high", "grade": "初三"},
                files=[
                    (
                        "file",
                        p1_docx.name,
                        p1_docx.read_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                ],
                timeout=args.timeout,
            )
        )
        assert p1_paper["status"] == "succeeded", p1_paper
        assert p1_paper["summary"]["question_count"] == 2, p1_paper
        p1_paper_job = envelope_data(get_json(f"{p2_base}/api/ai/v1/jobs/{p1_paper['job_id']}", timeout=args.timeout))
        assert p1_paper_job["job_type"] == "paper_parse", p1_paper_job
        assert p1_paper_job["result_url"].endswith(f"/parse/paper/{p1_paper['job_id']}/result"), p1_paper_job
        p1_paper_result = envelope_data(
            get_json(f"{p2_base}/api/ai/v1/parse/paper/{p1_paper['job_id']}/result", timeout=args.timeout)
        )
        assert p1_paper_result["questions"][0]["question_no"] == "1", p1_paper_result

    p1_score = envelope_data(
        post_multipart(
            f"{p2_base}/api/ai/v1/parse/score-excel",
            fields={},
            files=[
                (
                    "file",
                    scores_path.name,
                    scores_path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            ],
            timeout=args.timeout,
        )
    )
    assert p1_score["status"] == "succeeded", p1_score
    assert p1_score["summary"]["record_count"] >= 18, p1_score
    p1_score_job = envelope_data(get_json(f"{p2_base}/api/ai/v1/jobs/{p1_score['job_id']}", timeout=args.timeout))
    assert p1_score_job["job_type"] == "score_excel_parse", p1_score_job
    assert p1_score_job["result_url"].endswith(f"/parse/score-excel/{p1_score['job_id']}/result"), p1_score_job
    p1_score_result = envelope_data(
        get_json(f"{p2_base}/api/ai/v1/parse/score-excel/{p1_score['job_id']}/result", timeout=args.timeout)
    )
    assert len(p1_score_result["score_stats"]) == p1_score["summary"]["record_count"], p1_score_result

    exam = envelope_data(
        post_json(
            f"{p2_base}/exams",
            {
                "name": "Full-stack HTTP teacher smoke",
                "subject": "math",
                "grade": "senior_high",
                "class_ids": ["class_http_smoke"],
                "exam_date": "2026-07-30",
                "teacher_id": "teacher_http_smoke",
                "is_system_test": True,
            },
            timeout=args.timeout,
        )
    )
    exam_id = exam["exam_id"]

    paper_file = envelope_data(
        post_multipart(
            f"{p2_base}/exams/{exam_id}/files",
            fields={"file_type": "paper"},
            files=[("file", paper_path.name, paper_path.read_bytes(), "application/json")],
            timeout=args.timeout,
        )
    )["file"]
    assert len(paper_file["sha256"]) == 64 and paper_file["size_bytes"] > 0, paper_file
    score_file = envelope_data(
        post_multipart(
            f"{p2_base}/exams/{exam_id}/files",
            fields={"file_type": "score_excel"},
            files=[
                (
                    "file",
                    scores_path.name,
                    scores_path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            ],
            timeout=args.timeout,
        )
    )["file"]
    assert len(score_file["sha256"]) == 64 and score_file["size_bytes"] > 0, score_file

    parse_result = envelope_data(
        post_json(
            f"{p2_base}/exams/{exam_id}/parse",
            {
                "score_file_id": score_file["file_id"],
                "paper_file_id": paper_file["file_id"],
                "auto_tag_knowledge": True,
            },
            timeout=args.timeout,
        )
    )
    assert parse_result["status"] == "teacher_review", parse_result

    structure = envelope_data(get_json(f"{p2_base}/exams/{exam_id}/structure", timeout=args.timeout))
    assert len(structure["questions"]) == 18, structure
    first_question = structure["questions"][0]

    update_payload = envelope_data(
        put_json(
            f"{p2_base}/exams/{exam_id}/questions/{first_question['exam_question_id']}",
            {
                "question_no": first_question["question_no"],
                "stem_html": "<p>HTTP smoke teacher revised stem</p>",
                "question_type": first_question["question_type"] or "single_choice",
                "full_score": first_question["full_score"],
            },
            timeout=args.timeout,
        )
    )
    assert update_payload["updated"] is True, update_payload

    knowledge_ids = first_question["knowledge_point_ids"] or ["kp_http_smoke"]
    tag_payload = envelope_data(
        put_json(
            f"{p2_base}/exams/{exam_id}/questions/{first_question['exam_question_id']}/knowledge-tags",
            {"knowledge_point_ids": knowledge_ids, "comment": "HTTP smoke confirmation"},
            timeout=args.timeout,
        )
    )
    assert tag_payload["knowledge_point_ids"] == knowledge_ids, tag_payload

    analysis = envelope_data(get_json(f"{p2_base}/exams/{exam_id}/analysis", timeout=args.timeout))
    assert len(analysis["question_analysis"]) == 18, analysis
    assert analysis["knowledge_tag_coverage"] > 0, analysis
    assert "???" not in json.dumps(analysis.get("practice_recommendations", []), ensure_ascii=False), analysis.get(
        "practice_recommendations", []
    )
    assert any(
        item["stem_text"] == "HTTP smoke teacher revised stem"
        and item["teacher_review_status"] == "confirmed"
        for item in analysis["question_analysis"]
    ), analysis["question_analysis"][:2]

    diagnostics_run = envelope_data(
        post_json(
            f"{p2_base}/exams/{exam_id}/diagnostics/run",
            {
                "analysis_scope": "class",
                "class_id": "class_http_smoke",
                "include_teaching_suggestions": True,
                "include_question_recommendations": True,
            },
            timeout=args.timeout,
        )
    )
    diagnostic_id = diagnostics_run["diagnostic_id"]
    diagnostic = envelope_data(
        get_json(f"{p2_base}/exams/{exam_id}/diagnostics/{diagnostic_id}", timeout=args.timeout)
    )
    assert diagnostic["summary"]["question_count"] == 18, diagnostic
    assert diagnostic["summary"]["knowledge_tag_coverage"] > 0, diagnostic

    practice_pack = envelope_data(
        post_json(
            f"{p2_base}/exams/{exam_id}/practice-packs",
            {
                "diagnostic_id": diagnostic_id,
                "title": "HTTP smoke weakness practice pack",
                "target": "class",
                "target_ref_id": "class_http_smoke",
                "created_by": "teacher_http_smoke",
            },
            timeout=args.timeout,
        )
    )["practice_pack"]
    assert practice_pack["practice_pack_id"].startswith("pack_"), practice_pack
    assert practice_pack["question_ids"], practice_pack

    practice_packs = envelope_data(get_json(f"{p2_base}/exams/{exam_id}/practice-packs", timeout=args.timeout))
    assert any(item["practice_pack_id"] == practice_pack["practice_pack_id"] for item in practice_packs["items"])

    lesson_plan = envelope_data(
        post_json(
            f"{p2_base}/exams/{exam_id}/lesson-plans",
            {
                "diagnostic_id": diagnostic_id,
                "template_id": "tpl_school_math_review_v1",
                "sections": ["exam_summary", "high_loss_questions", "weakness_summary"],
            },
            timeout=args.timeout,
        )
    )
    download_url = lesson_plan["file"]["download_url"]
    assert download_url.endswith("/download"), lesson_plan

    lesson_plans = envelope_data(get_json(f"{p2_base}/exams/{exam_id}/lesson-plans", timeout=args.timeout))
    assert any(item["download_url"] == download_url for item in lesson_plans["items"]), lesson_plans

    lesson_bytes = get_bytes(abs_url(p2_base, download_url), timeout=args.timeout)
    assert len(lesson_bytes) > 10_000, len(lesson_bytes)

    exam_summary = envelope_data(get_json(f"{p2_base}/exams/{exam_id}", timeout=args.timeout))
    assert exam_summary["question_count"] == 18, exam_summary
    assert exam_summary["lesson_plan_count"] >= 1, exam_summary
    assert exam_summary["practice_pack_count"] >= 1, exam_summary
    assert exam_summary["latest_lesson_plan"]["download_url"].endswith("/download"), exam_summary
    assert exam_summary["latest_practice_pack"]["practice_pack_id"].startswith("pack_"), exam_summary
    default_exam_list = envelope_data(get_json(f"{p2_base}/exams", timeout=args.timeout))
    assert all(item["exam_id"] != exam_id for item in default_exam_list["items"]), default_exam_list
    system_exam_list = envelope_data(get_json(f"{p2_base}/exams?include_system=true", timeout=args.timeout))
    assert any(item["exam_id"] == exam_id and item["is_system_test"] is True for item in system_exam_list["items"]), (
        system_exam_list
    )
    readiness_after_system_test = envelope_data(get_json(f"{p2_base}/api/demo/readiness", timeout=args.timeout))
    assert readiness_after_system_test["facts"]["exam_count"] == initial_customer_exam_count, readiness_after_system_test
    assert readiness_after_system_test["facts"]["lesson_plan_count"] == initial_customer_lesson_plan_count, (
        readiness_after_system_test
    )
    assert readiness_after_system_test["facts"]["practice_pack_count"] == initial_customer_practice_pack_count, (
        readiness_after_system_test
    )
    assert readiness_after_system_test["facts"]["system_test_exam_count"] >= initial_system_test_exam_count + 1, (
        readiness_after_system_test
    )

    audit_log = envelope_data(get_json(f"{p2_base}/audit-logs?limit=120", timeout=args.timeout))
    audit_events = {item["event"] for item in audit_log["items"]}
    for expected_event in {
        "exam_parse_succeeded",
        "knowledge_tags_confirmed",
        "diagnostic_run",
        "ai_generated_question_reviewed",
        "lesson_plan_downloaded",
    }:
        assert expected_event in audit_events, audit_events

    frontend_url = f"{frontend_base}/?apiBase={parse.quote(p2_base, safe='')}"
    frontend_html = get_text(frontend_url, timeout=args.timeout)
    assert "root" in frontend_html and "校园智能学情系统" in frontend_html, frontend_html[:500]

    print("campus-system full-stack HTTP smoke passed")
    print(f"p2_url={p2_base}")
    print(f"p3_url={p3_base}")
    print(f"frontend_url={frontend_url}")
    print(f"readiness={readiness['overall_status']}")
    print(f"papers={readiness['facts']['paper_count']}")
    print(f"questions={readiness['facts']['question_count']}")
    print(f"knowledge_points={readiness['facts']['knowledge_point_count']}")
    print(f"teacher_session_role={teacher_session['actor']['actor_role']}")
    print(f"student_session_role={student_session['actor']['actor_role']}")
    print(f"identity_directory_source={teacher_session['identity_directory']['source']}")
    print(f"identity_directory_users={teacher_session['identity_directory']['user_count']}")
    print(f"p3_search_items={len(p3_search['items'])}")
    print(f"student_practice_items={len(student_recommendations['items'])}")
    print(f"student_progress_answers={student_progress['answer_count']}")
    print(f"student_history_items={len(student_history['items'])}")
    print(f"student_report_level={student_report['summary']['report_level']}")
    print(f"student_wrong_question={student_wrong_confirm['status']}")
    print(f"student_variants={len(variants['items'])}")
    print(f"generated_review={generated_review['audit_status']}")
    print(f"audit_events={len(audit_log['items'])}")
    print(f"teacher_exam_id={exam_id}")
    print(f"teacher_questions={len(analysis['question_analysis'])}")
    print(f"teacher_knowledge_tag_coverage={analysis['knowledge_tag_coverage']}")
    print(f"teacher_practice_pack_questions={len(practice_pack['question_ids'])}")
    print(f"teacher_lesson_bytes={len(lesson_bytes)}")


def get_json(url: str, *, headers: dict[str, str] | None = None, timeout: float = 12) -> dict:
    return json.loads(get_bytes(url, headers=headers, timeout=timeout).decode("utf-8"))


def get_json_or_none(url: str, *, headers: dict[str, str] | None = None, timeout: float = 12) -> dict | None:
    try:
        return get_json(url, headers=headers, timeout=timeout)
    except error.HTTPError as exc:
        if exc.code == 404:
            return None
        raise


def ensure_publish_smoke_generated_question(
    *,
    p2_base: str,
    p3_base: str,
    source_question_id: str,
    variant: dict,
    model_name: str,
    timeout: float,
) -> dict:
    existing = get_generated_question_or_none(
        p3_base=p3_base,
        generated_question_id=PUBLISH_GENERATED_SMOKE_ID,
        timeout=timeout,
    )
    if existing:
        if existing["audit_status"] == "approved":
            assert existing.get("bank_question_id"), existing
            return {
                "generated_question_id": existing["generated_question_id"],
                "audit_status": existing["audit_status"],
                "bank_question_id": existing["bank_question_id"],
            }
        assert existing["audit_status"] == "pending_review", existing
    else:
        submit_payload = envelope_data(
            post_json(
                f"{p2_base}/ai-generated-questions",
                generated_question_import_payload(
                    source_question_id=source_question_id,
                    model_name=model_name,
                    items=[generated_question_item(variant, PUBLISH_GENERATED_SMOKE_ID)],
                ),
                timeout=timeout,
            )
        )
        assert submit_payload["saved_count"] == 1, submit_payload
        assert submit_payload["items"][0]["generated_question_id"] == PUBLISH_GENERATED_SMOKE_ID, submit_payload

    return envelope_data(
        put_json(
            f"{p2_base}/ai-generated-questions/{PUBLISH_GENERATED_SMOKE_ID}/review",
            {
                "decision": "approved",
                "reviewer_id": "teacher_http_smoke",
                "review_comment": "One-time HTTP smoke publish check.",
                "publish_to_bank": True,
            },
            timeout=timeout,
        )
    )


def get_generated_question_or_none(*, p3_base: str, generated_question_id: str, timeout: float) -> dict | None:
    response = get_json_or_none(
        f"{p3_base}/api/resource/v1/generated-questions/{generated_question_id}",
        headers={"X-Service-Id": "p2-service"},
        timeout=timeout,
    )
    if response is None:
        return None
    return envelope_data(response)


def generated_question_import_payload(*, source_question_id: str, model_name: str, items: list[dict]) -> dict:
    return {
        "source_question_id": source_question_id,
        "knowledge_point_version": "2026.1",
        "model_name": model_name,
        "prompt_version": "full_stack_http_smoke.v1",
        "items": items,
    }


def generated_question_item(variant: dict, generated_question_id: str) -> dict:
    return {
        **variant,
        "generated_question_id": generated_question_id,
        "question_type": "single_choice",
        "knowledge_point_ids": ["kp_math_8_triangle_side_relation"],
    }


def write_smoke_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("一、选择题")
    document.add_paragraph("1．（5分）若 |x|=2，则 x 的值可能是（　　）")
    document.add_paragraph("A．-2 B．0 C．1 D．3")
    document.add_paragraph("二、解答题")
    document.add_paragraph("2．（5分）解方程 x^2=4，并说明理由。")
    document.save(path)


def get_text(url: str, *, headers: dict[str, str] | None = None, timeout: float = 12) -> str:
    return get_bytes(url, headers=headers, timeout=timeout).decode("utf-8", errors="replace")


def get_bytes(url: str, *, headers: dict[str, str] | None = None, timeout: float = 12) -> bytes:
    req = request.Request(url, headers=headers or {}, method="GET")
    with request.urlopen(req, timeout=timeout) as response:
        return response.read()


def expect_http_status(url: str, expected_status: int, *, timeout: float = 12) -> None:
    req = request.Request(url, method="GET")
    try:
        with request.urlopen(req, timeout=timeout) as response:
            status = response.status
    except error.HTTPError as exc:
        status = exc.code
    assert status == expected_status, {"url": url, "expected": expected_status, "actual": status}


def post_json(
    url: str,
    payload: dict,
    *,
    headers: dict[str, str] | None = None,
    timeout: float = 12,
) -> dict:
    request_headers = {"Content-Type": "application/json", **(headers or {})}
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = request.Request(url, data=body, headers=request_headers, method="POST")
    with request.urlopen(req, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def put_json(
    url: str,
    payload: dict,
    *,
    headers: dict[str, str] | None = None,
    timeout: float = 12,
) -> dict:
    request_headers = {"Content-Type": "application/json", **(headers or {})}
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = request.Request(url, data=body, headers=request_headers, method="PUT")
    with request.urlopen(req, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def post_multipart(
    url: str,
    *,
    fields: dict[str, str],
    files: list[tuple[str, str, bytes, str]],
    headers: dict[str, str] | None = None,
    timeout: float = 12,
) -> dict:
    boundary = f"----CampusSmoke{uuid4().hex}"
    parts: list[bytes] = []
    for name, value in fields.items():
        parts.append(
            (
                f"--{boundary}\r\n"
                f"Content-Disposition: form-data; name=\"{name}\"\r\n\r\n"
                f"{value}\r\n"
            ).encode("utf-8")
        )
    for field_name, filename, content, content_type in files:
        safe_filename = filename.replace("\\", "_").replace('"', "_")
        parts.append(
            (
                f"--{boundary}\r\n"
                f"Content-Disposition: form-data; name=\"{field_name}\"; filename=\"{safe_filename}\"\r\n"
                f"Content-Type: {content_type}\r\n\r\n"
            ).encode("utf-8")
            + content
            + b"\r\n"
        )
    parts.append(f"--{boundary}--\r\n".encode("utf-8"))
    body = b"".join(parts)
    request_headers = {"Content-Type": f"multipart/form-data; boundary={boundary}", **(headers or {})}
    req = request.Request(url, data=body, headers=request_headers, method="POST")
    with request.urlopen(req, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def abs_url(base_url: str, maybe_relative_url: str) -> str:
    if maybe_relative_url.startswith(("http://", "https://")):
        return maybe_relative_url
    return f"{base_url}{maybe_relative_url if maybe_relative_url.startswith('/') else '/' + maybe_relative_url}"


def envelope_data(payload: dict) -> dict:
    assert payload.get("code") in {"OK", "ok"}, payload
    assert isinstance(payload.get("data"), dict), payload
    return payload["data"]


if __name__ == "__main__":
    main()
